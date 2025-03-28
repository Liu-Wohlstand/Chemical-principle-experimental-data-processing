import matplotlib.pyplot as plt
import numpy as np
import csv
import matplotlib
from scipy.optimize import curve_fit
from matplotlib.font_manager import FontProperties
from docx import Document

from docx.oxml.ns import qn  # 用于设置中文字体


times_new_roman = FontProperties(family='Times New Roman')
matplotlib.rcParams['font.family'] = 'SimSun'  # 'SimSun' 是宋体
matplotlib.rcParams['axes.unicode_minus'] = False

# -----以上请勿改动--------------------------------------------------------------
name = 'lihua'#一般将实验者姓名作为.csv的命名，如:李华做了实验则普通管文件命名为lihua.csv，强化管命名为lihua2.csv
# ------以下请勿改动-------------------------------------------------------------
geshi = '.csv'
qianghuaguan = name + geshi
putongguan = name + '2' + geshi
mingming = name+'数据处理'+ '.docx'
di = 19.25 * 10 ** -3  # 内管管内径 m
Li = 1  # 传热管测量段实际测量长度 m

doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

ls = []
with open(qianghuaguan, 'r', encoding='utf-8') as file:
    reader = csv.reader(file)
    for row in reader:
        str_data = ','.join(row)
        str_data = str_data.replace('\ufeff', '')
        result = str_data.split(',')
        float_row = [float(item) for item in result]
        ls.append(float_row)
#  ls1: ls[0]压差计 ls[1]加热电压 ls[2]壁面温度 ls[3]进液温度 ls[4]出液温度
ls_p = []  # 压差计
ls_bi_T = []  # 壁面温度
ls_jin_T = []  # 进液温度
ls_chu_T = []  # 出液温度
for i in ls:
    ls_p.append(i[0])
    ls_bi_T.append(i[2])
    ls_jin_T.append(i[3])
    ls_chu_T.append(i[4])

ls2 = []
with open(putongguan, 'r', encoding='utf-8') as file:
    reader = csv.reader(file)
    for row in reader:
        str_data = ','.join(row)
        str_data = str_data.replace('\ufeff', '')
        result = str_data.split(',')
        float_row = [float(item) for item in result]
        ls2.append(float_row)

#  ls1: ls[0]压差计 ls[1]加热电压 ls[2]壁面温度 ls[3]进液温度 ls[4]出液温度 ls[5]流量计 ls[6]平均温度 ls[7]
ls_p2 = []  # 压差计
ls_bi_T2 = []  # 壁面温度
ls_jin_T2 = []  # 进液温度
ls_chu_T2 = []  # 出液温度

for i in ls2:
    ls_p2.append(i[0])
    ls_bi_T2.append(i[2])
    ls_jin_T2.append(i[3])
    ls_chu_T2.append(i[4])


def chazhi(T, dic):
    if T in dic:
        return dic[T]
    sorted_keys = sorted(dic.keys())
    lower_key = None
    upper_key = None
    for key in sorted_keys:
        if key < T:
            lower_key = key
        else:
            upper_key = key
            break
    if lower_key is None or upper_key is None:
        print(f"插值无法计算，{T}超出字典{dic}键的范围")
    lower_value = dic[lower_key]
    upper_value = dic[upper_key]
    interpolation = (T - lower_key) * (upper_value - lower_value) / (upper_key - lower_key) + lower_value
    return interpolation


'''# 以下所有数据均用水的物性
Pr_dict = {30: 5.42, 40: 4.32, 50: 3.54, 60: 2.98, 70: 2.54, 80: 2.22, 90: 1.96}
# 黏度需要乘10^5使用
viscosity_dict = {30: 80.07,
                  40: 65.6,
                  50: 54.94,
                  60: 46.88,
                  70: 40.61,
                  80: 35.65,
                  90: 31.65}

density_dict = {30: 995.7,
                40: 992.2,
                50: 988.1,
                60: 983.2,
                70: 977.8,
                80: 971.8,
                90: 965.3}
# 热导率需乘10^2使用
thermal_conductivity_dict = {30: 61.76,
                             40: 63.38,
                             50: 64.78,
                             60: 65.94,
                             70: 66.76,
                             80: 67.45,
                             90: 68.04}
specific_heat_dict = {30: 4.174,
                      40: 4.174,
                      50: 4.174,
                      60: 4.178,
                      70: 4.187,
                      80: 4.195,
                      90: 4.208}'''
# 以下所有数据均用干空气的物性
Pr_dict = {20: 0.703, 30: 0.701, 40: 0.699, 50: 0.698, 60: 0.696, 70: 0.694, 80: 0.692, 90: 0.690}
# 黏度需要乘10^5使用
viscosity_dict = {20: 1.81, 30: 1.86, 40: 1.91, 50: 1.96, 60: 2.01, 70: 2.06, 80: 2.11, 90: 2.15}

density_dict = {20: 1.205, 30: 1.165, 40: 1.128, 50: 1.093, 60: 1.060, 70: 1.029, 80: 1, 90: 0.972}
# 热导率需乘10^2使用
thermal_conductivity_dict = {20: 2.593,
                             30: 2.675,
                             40: 2.756,
                             50: 2.826,
                             60: 2.896,
                             70: 2.966,
                             80: 3.047,
                             90: 3.128}
specific_heat_dict = {20: 1.005,
                      30: 1.005,
                      40: 1.005,
                      50: 1.005,
                      60: 1.005,
                      70: 1.009,
                      80: 1.009,
                      90: 1.009}

Si = np.pi * di * Li
print(f"管路换热面积为{Si} m^2")


def shujuchuli(tw, ti1, ti2, p):
    delta_tm = ((tw - ti1) - (tw - ti2)) / np.log((tw - ti1) / (tw - ti2))
    cpi = chazhi((ti2 + ti1) / 2, specific_heat_dict)  # 冷流体的定压比热
    roui = chazhi((ti2 + ti1) / 2, density_dict)  # 冷流体的密度
    miui = chazhi((ti2 + ti1) / 2, viscosity_dict)
    lambdai = chazhi((ti2 + ti1) / 2, thermal_conductivity_dict)
    Pri = chazhi((ti2 + ti1) / 2, Pr_dict)  # Pr随温度有变化为什么默认不变？因为空气的变化太小了……但是水的变化很大
    rou_t0 = chazhi(ti1, density_dict)

    Vt0 = 23.80 * (p / rou_t0) ** 0.5
    Vi = Vt0 * ((273 + (ti2 + ti1) / 2) / (273 + ti1))
    Wi = Vi * roui / 3600  # 此处Vi单位为平方米每小时
    Qi = Wi * cpi * 1000 * (ti2 - ti1)
    ai = Qi / (delta_tm * Si)
    A = 0.25 * np.pi * di ** 2
    ui = Vi / 3600 / A
    Nui = ai * di / (lambdai * 10 ** -2)
    Rei = ui * di * roui / (miui * 10 ** -5)
    print(f'压差为{p} kPa下的数据如下：\n'
          f'管内定性（平均）温度tm={(ti2 + ti1) / 2} ℃\n'
          f'管内温度差={ti2 - ti1} ℃\n'
          f'对数平均温差Δtmi={delta_tm:.2f} ℃\n'
          f'空气平均密度ρi={roui:.3f} kg/m^3\n'
          f'入口处空气密度ρt0={rou_t0:.3f} kg/m^3\n'
          f'空气热导率λi={lambdai:.3f} *10^-2W/(m·℃)\n'
          f'空气比热ci={cpi:.3f} kJ/(kg·℃)\n'
          f'空气黏度μ={miui:.2f} *10^-5Pa·s\n'
          f'普朗特数Pr={Pri:.3f}\n'
          f'流量{Vt0:.2f} m^3/h\n'
          f'校正流量{Vi:.2f} m^3/h\n'
          f'质量流量{Wi:.4f} kg/s\n'
          f'传热速率Qi={Qi}\n'
          f'传热系数ai={ai:.4f} W/(m^2·℃)\n'
          f'管路面积{A}m^2\n'
          f'流速ui={ui:.3f} m/s\n'
          f'Rei={Rei:.2f}\n'
          f'Nui={Nui:.2f}'
          f'\n')
    paragraph = doc.add_paragraph(f'压差为{p} kPa下的数据如下：\n'
          f'管内定性（平均）温度tm={(ti2 + ti1) / 2} ℃\n'
          f'管内温度差={ti2 - ti1} ℃\n'
          f'对数平均温差Δtmi={delta_tm:.2f} ℃\n'
          f'空气平均密度ρi={roui:.3f} kg/m^3\n'
          f'入口处空气密度ρt0={rou_t0:.3f} kg/m^3\n'
          f'空气热导率λi={lambdai:.3f} *10^-2W/(m·℃)\n'
          f'空气比热ci={cpi:.3f} kJ/(kg·℃)\n'
          f'空气黏度μ={miui:.2f} *10^-5Pa·s\n'
          f'普朗特数Pr={Pri:.3f}\n'
          f'流量{Vt0:.2f} m^3/h\n'
          f'校正流量{Vi:.2f} m^3/h\n'
          f'质量流量{Wi:.4f} kg/s\n'
          f'传热速率Qi={Qi}\n'
          f'传热系数ai={ai:.4f} W/(m^2·℃)\n'
          f'管路面积{A}m^2\n'
          f'流速ui={ui:.3f} m/s\n'
          f'Rei={Rei:.2f}\n'
          f'Nui={Nui:.2f}'
          f'\n')


    run = paragraph.add_run()

    return Rei, Nui, Pri


ls_Nu1 = []
ls_Re1 = []
ls_Nu2 = []
ls_Re2 = []
ls_Pr1 = []
ls_Pr2 = []
print("实验一数据处理如下：")
for i in range(len(ls_chu_T)):
    Rei, Nui, Pri = shujuchuli(ls_bi_T[i], ls_jin_T[i], ls_chu_T[i], ls_p[i])
    ls_Nu1.append(Nui)
    ls_Re1.append(Rei)
    ls_Pr1.append(Pri)
print("实验二数据处理如下：")
for i in range(len(ls_chu_T2)):
    Rei2, Nui2, Pri2 = shujuchuli(ls_bi_T2[i], ls_jin_T2[i], ls_chu_T2[i], ls_p2[i])
    ls_Nu2.append(Nui2)
    ls_Pr2.append(Pri2)
    ls_Re2.append(Rei2)


def model_func(x, a, b):
    return a * np.power(x, b)


def fit_power_law(x_data, y_data):
    params, covariance = curve_fit(model_func, x_data, y_data, p0=[1, 1], maxfev=10000)
    a, b = params
    predicted_y = model_func(x_data, a, b)
    correlation_coefficient = np.corrcoef(y_data, predicted_y)[0, 1]
    return a, b, correlation_coefficient


def zhixiannihe(x, y):
    coefficients = np.polyfit(x, y, 1)
    fit_function = np.poly1d(coefficients)
    x_plot = np.linspace(min(x), max(x))
    y_plot = fit_function(x_plot)
    y_pred = fit_function(x)
    ss_res = np.sum((y - y_pred) ** 2)  # 残差平方和
    ss_tot = np.sum((y - np.mean(y)) ** 2)  # 总平方和
    r_squared = 1 - (ss_res / ss_tot)
    return x_plot, y_plot, coefficients[0], coefficients[1], r_squared


x1, y1, k1, b1, r1 = zhixiannihe(np.log10(ls_Re1), np.log10(np.array(ls_Nu1) / np.array(ls_Pr1) ** 0.4))
x2, y2, k2, b2, r2 = zhixiannihe(np.log10(ls_Re2), np.log10(np.array(ls_Nu2) / np.array(ls_Pr2) ** 0.4))
plt.plot(x2, y2, label='普通管')
plt.plot(x1, y1, label='强化管', color='green')
plt.scatter(np.log10(ls_Re1), np.log10(np.array(ls_Nu1) / np.array(ls_Pr1) ** 0.4), color='green')
plt.scatter(np.log10(ls_Re2), np.log10(np.array(ls_Nu2) / np.array(ls_Pr2) ** 0.4))
plt.xlabel('lg(Re)')
plt.ylabel('lg(Nu/Pr^0.4)')
plt.title('准数关联图')
plt.legend()
plt.savefig(name + '准数关联图' + '.svg', dpi=600, facecolor='w', edgecolor='w',
            orientation='portrait', format='svg',
            transparent=False, bbox_inches=None, pad_inches=0.1,
            metadata=None)
plt.savefig(name + '准数关联图' + '.png')

plt.show()
print(f'k1={k1},b1={b1},r1={r1}')
print(f'k2={k2},b2={b2},r2={r2}')

a1, b1, correlation_coefficient1 = fit_power_law(ls_Re1, np.array(ls_Nu1) / np.array(ls_Pr1) ** 0.4)
print(f'第一组实验的斜率A={a1},幂b={b1},相关系数R^2={correlation_coefficient1}')
a2, b2, correlation_coefficient2 = fit_power_law(ls_Re2, np.array(ls_Nu2) / np.array(ls_Pr2) ** 0.4)
print(f'第二组实验的斜率A={a2},幂b={b2},相关系数R^2={correlation_coefficient2}')

# 绘制数据点和拟合曲线
plt.scatter(ls_Re1, np.array(ls_Nu1) / np.array(ls_Pr1) ** 0.4)
plt.plot(ls_Re1, model_func(ls_Re1, a1, b1), label='强化管')
plt.scatter(ls_Re2, np.array(ls_Nu2) / np.array(ls_Pr2) ** 0.4)
plt.plot(ls_Re2, model_func(ls_Re2, a2, b2), label='普通管')
plt.xlabel('Re')
plt.ylabel('Nu/Pr^0.4')
plt.title('准数关联图')
plt.legend()
plt.show()
"""doc.add_picture(name+'准数关联图'+'.svg')"""
doc.add_picture(name+'准数关联图'+'.png')
doc.save(mingming)
