"""
测定水流量40 L/h下的Δp/Z-u关系曲线
以空气流量为8 m^3/h 和 11 m^3/h，并以混合气中氨组成为0.03摩尔比 水流量30 L/h进行两次吸收实验
"""
import matplotlib.pyplot as plt
import openpyxl
import numpy as np
from scipy.interpolate import UnivariateSpline
from matplotlib.font_manager import FontProperties
import matplotlib
matplotlib.rcParams['font.family'] = 'Arial'
from docx import Document
from docx.oxml.ns import qn  # 用于设置中文字体
from matplotlib.ticker import FixedLocator
doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
# -----以上请勿改动--------------------------------------------------------------
name = 'lihua'#一般将实验者姓名作为.csv的命名，如:李华做了实验则干燥无喷淋液文件命名为lihua干燥.xlsx，水位吸收液命名为lihua水.xlsx
# ------以下请勿改动-------------------------------------------------------------

geshi = '.xlsx'
ganzao = name + '干燥' + geshi
shui = name + '水' + geshi
mingming = name + '吸收实验数据处理报告' + '.docx'
def duqu(name):
    wb = openpyxl.load_workbook(name)
    sheet = wb.active
    column_data = []
    column_data2 = []
    for row in sheet.iter_rows(min_row=2, max_col=2, values_only=True):  # min_row=2 跳过表头
        # 将每行的数据添加到对应的列数组中
        column_data.append(row[0])
        column_data2.append(row[1])


    return np.array(column_data),np.array(column_data2)


D = 0.075  # 填料塔内径 m
Z = 0.4  # 填料层高度 m

def zhixiannihe(x, y):
    coefficients = np.polyfit(x, y, 1)
    fit_function = np.poly1d(coefficients)
    x_plot = np.linspace(min(x), max(x), 100)
    y_plot = fit_function(x_plot)
    return x_plot, y_plot

V,delta_p = duqu(ganzao)
V2,delta_p2 = duqu(shui)
delta_p_Z = delta_p/Z*98.0665
delta_p2_Z = delta_p2/Z*98.0665
u1=V/(np.pi*3600*0.25*D**2)
u2=V2/(np.pi*3600*0.25*D**2)
u1_nihe,delta_p_Z_nihe=zhixiannihe(u1,delta_p_Z)

print(f'下面是干填料塔数据处理：')
paragraph = doc.add_paragraph(f'下面是干填料塔数据处理：')
for i in range(len(u1)):
    print(f'第{i+1}组，压强为{delta_p[i]}\n'
          f'p={delta_p[i]*98.0665:.2f} Pa\n'
          f'Δp/Z={delta_p_Z[i]:.2f}\n'
          f'u={u1[i]:.2f}\n'
            )
    paragraph2 = doc.add_paragraph(f'第{i+1}组，压强为{delta_p[i]}\n'
          f'p={delta_p[i]*98.0665:.2f} Pa\n'
          f'Δp/Z={delta_p_Z[i]:.2f}\n'
          f'u={u1[i]:.2f}\n'
            )
print(f'下面是湿填料塔数据处理：')
paragraph3 = doc.add_paragraph(f'下面是湿填料塔数据处理：')
for i in range(len(u2)):
    print(f'第{i+1}组，压强为{delta_p2[i]}\n'
          f'p={delta_p2[i]*98.0665:.2f} Pa\n'
          f'Δp/Z={delta_p2_Z[i]:.2f}\n'
          f'u={u2[i]:.2f}\n'
            )
    paragraph4 = doc.add_paragraph(f'第{i+1}组，压强为{delta_p2[i]}\n'
          f'p={delta_p2[i]*98.0665:.2f} Pa\n'
          f'Δp/Z={delta_p2_Z[i]:.2f}\n'
          f'u={u2[i]:.2f}\n')
#plt.loglog(u1_nihe, delta_p_Z_nihe, base=10)
plt.scatter(u1,delta_p_Z,marker='x')
plt.loglog(u1, delta_p_Z, base=10)
plt.loglog(u2, delta_p2_Z, base=10)
plt.scatter(u2,delta_p2_Z,marker='x')
plt.xlabel('u (m/s)')
plt.ylabel('Δp/Z (Pa/m)')
plt.title('Δp/Z-u')
ax = plt.gca()

ax.grid(which='both', color='gray', linewidth=1)
plt.show()
plt.plot(u1_nihe,delta_p_Z_nihe)
plt.scatter(u1,delta_p_Z)
plt.show()
M_acid = 98
'''def shujuchuli():
    Y2 = 2 * M_acid * V_acid / ((V_liangqiguan * T0 / T_liangqiguan) / 22.4)
    omega = np.pi * D ** 2  # 填料塔截面积，m2
    deltaY1=Y1-Y1_
    deltaY2=Y2-Y2_
    delta_Ym = (deltaY1 - deltaY2) / np.log(deltaY1 / deltaY2)  # 所测填料层两端面上气相推动力的平均值
    N_OG = (Y1 - Y2) / delta_Ym  # 气相总传质单元数，无因次
    H_OG = Z / N_OG  # 气相总传质单元高度，m
    V=#空气的摩尔流率，kmol（B）/ h
    K_Ya = V / (H_OG * omega)  # 气相总体积吸收系数，kmol /（m3 · h）
    psi_A = (Y1 - Y2) / Y1  # 混合气中氨被吸收的百分率（吸收率），无因次
    # 操作条件下喷淋密度的计算
    U = qv / omega  # qv为流量 m^3/h，omega为截面积 m^2


ls_P = []
ls_delta_P = np.array()
# 逐差法求压降
for i in range(len(ls_P)):
    if i > 0:
        delta_P = ls_P[i] - ls_P[i - 1]
    ls_delta_P = np.append(delta_P)
'''
doc.save(mingming)
